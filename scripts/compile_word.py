from supabase import create_client
import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import re
from datetime import datetime

now = datetime.now()

# Função para impedir a divisão da tabela entre páginas
def prevent_table_split(table):
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

# Detalhes de conexão com Supabase
url = "https://rhzbwahkrnsrmpgozhyq.supabase.co"
key = os.environ.get('SUPABASE_API_KEY')
supabase = create_client(url, key)

# Buscar dados da tabela 'Notícias'
response = supabase.table('compile_adhoc').select('*').execute()
df = pd.DataFrame(response.data)

# Checar se há notícias
if df.empty:
    with open('flag_noticias.txt', 'w') as flag_file:
        flag_file.write('nao')
else:
    with open('flag_noticias.txt', 'w') as flag_file:
        flag_file.write('sim')
    
    doc = Document()

    # Extrair texto TLDR
    df['TLDR_text'] = df['TLDR'].str.extract(r'TLDR:\s*([\s\S]*)', flags=re.DOTALL)

    # Iterar por cada linha do dataframe e adicionar ao documento
    for index, row in df.iterrows():
        # Adicionar uma tabela com uma única célula para encapsular todo o bloco de informação
        block_table = doc.add_table(rows=1, cols=1)
        block_cell = block_table.cell(0, 0)
        block_paragraph = block_cell.paragraphs[0]

        # Adicionar o título
        run = block_paragraph.add_run(row['título'] + '\n\n')
        run.bold = True
        block_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Detalhes da notícia como fonte, data de publicação, etc.
        details_text = f"Fonte: {row['fonte']}\nData de publicação: {pd.to_datetime(row['data_de_publicação']).strftime('%Y-%m-%d')}\nGUID: {row['guid']}\nValor publicitário: {row['valor_publicitário']} €\n\n"
        block_paragraph.add_run(details_text)

        # Verificar se 'TLDR_text' não está vazio e adicionar
        if pd.notnull(row['TLDR_text']) and row['TLDR_text'].strip():
            block_paragraph.add_run(row['TLDR_text'] + '\n\n')

        for run in block_paragraph.runs:
            run.font.size = Pt(10)

        prevent_table_split(block_table)

    # Nome do arquivo e salvando
    file_name = f'relatorio_media_{now.strftime("%Y_%m_%d_%H_%M")}.docx'
    output_path = os.path.join('backups/', file_name)
    doc.save(output_path)

    output_path_email = 'relatorio_media_semanal.docx'
    doc.save(output_path_email)